"""
Parser for DOCX files using python-docx.
"""
import docx
import logging

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path):
    """
    Extracts text from a DOCX file, including paragraphs and tables.
    """
    text_parts = []
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_path}: {e}")
        return ""
