"""End-to-end cover of the upload handler, from Telegram document to stored resume."""

from __future__ import annotations

from pathlib import Path
from typing import Any

import docx
import pytest
from tests.conftest import RESUME_TEXT, FakeContext, FakeMessage, FakeUpdate, FakeUser

from ats_bot.config import Settings
from ats_bot.db import repository
from ats_bot.handlers.common import STATE_AWAITING_JD, SessionKey
from ats_bot.handlers.upload import handle_resume_upload


def _write_docx(path: Path, text: str) -> Path:
    document = docx.Document()
    for line in text.split("\n"):
        document.add_paragraph(line)
    document.save(str(path))
    return path


class FakeDocument:
    """Stands in for ``telegram.Document``; serves a file from local disk."""

    def __init__(self, source: Path, file_name: str | None = None, size: int | None = None) -> None:
        self._source = source
        self.file_name = file_name if file_name is not None else source.name
        self.file_size = size if size is not None else source.stat().st_size
        self.file_id = "fake-file-id"

    async def get_file(self, **kwargs: Any) -> FakeDocument:
        return self

    async def download_to_drive(self, custom_path: Any, **kwargs: Any) -> None:
        Path(custom_path).write_bytes(self._source.read_bytes())


@pytest.fixture
def resume_docx(tmp_path: Path) -> Path:
    """A real .docx file containing the sample resume."""
    source = tmp_path / "source"
    source.mkdir(parents=True, exist_ok=True)
    return _write_docx(source / "cv.docx", RESUME_TEXT)


def _upload(document: Any, user_id: int = 42) -> tuple[FakeUpdate, FakeMessage]:
    message = FakeMessage()
    message.document = document
    return FakeUpdate(message=message, user=FakeUser(user_id)), message


class TestSuccessfulUpload:
    async def test_stores_the_resume_and_prompts_for_a_posting(
        self, db: Settings, context: FakeContext, resume_docx: Path
    ) -> None:
        update, message = _upload(FakeDocument(resume_docx))

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        stored = repository.get_all_resumes(42, settings=db)
        assert len(stored) == 1
        assert stored[0].file_name == "cv.docx"
        assert stored[0].file_type == "docx"
        assert "Backend engineer" in stored[0].extracted_text

        output = message.all_output
        assert "Resume received" in output
        assert "job description" in output

    async def test_sets_the_session_state(
        self, db: Settings, context: FakeContext, resume_docx: Path
    ) -> None:
        update, _ = _upload(FakeDocument(resume_docx))

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        assert context.user_data[SessionKey.STATE] == STATE_AWAITING_JD
        assert context.user_data[SessionKey.PENDING_RESUME_IDS]
        assert context.user_data[SessionKey.SESSION_CLEARED] is False

    async def test_file_lands_in_the_user_directory(
        self, db: Settings, settings: Settings, context: FakeContext, resume_docx: Path
    ) -> None:
        update, _ = _upload(FakeDocument(resume_docx))

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        stored = repository.get_all_resumes(42, settings=db)[0]
        assert Path(stored.file_path).is_file()
        assert Path(stored.file_path).parent == settings.upload_dir / "42"

    async def test_a_second_similar_upload_is_flagged_as_a_revision(
        self, db: Settings, context: FakeContext, tmp_path: Path
    ) -> None:
        first = _write_docx(tmp_path / "v1.docx", RESUME_TEXT)
        second = _write_docx(tmp_path / "v2.docx", RESUME_TEXT + "\nAdded a line about Kubernetes.")

        update_one, _ = _upload(FakeDocument(first))
        await handle_resume_upload(update_one, context)  # type: ignore[arg-type]

        update_two, message_two = _upload(FakeDocument(second))
        await handle_resume_upload(update_two, context)  # type: ignore[arg-type]

        assert "New version received" in message_two.all_output

    async def test_a_hostile_filename_cannot_escape_the_upload_directory(
        self, db: Settings, settings: Settings, context: FakeContext, resume_docx: Path
    ) -> None:
        update, _ = _upload(FakeDocument(resume_docx, file_name="../../../evil.docx"))

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        stored = repository.get_all_resumes(42, settings=db)[0]
        assert settings.upload_dir in Path(stored.file_path).parents


class TestRejections:
    async def test_unsupported_extension(
        self, db: Settings, context: FakeContext, tmp_path: Path
    ) -> None:
        text_file = tmp_path / "notes.txt"
        text_file.write_text("hello")
        update, message = _upload(FakeDocument(text_file))

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        assert "cannot read" in message.last_reply
        assert repository.get_all_resumes(42, settings=db) == []

    async def test_oversized_file(
        self, db: Settings, settings: Settings, context: FakeContext, resume_docx: Path
    ) -> None:
        oversized = FakeDocument(resume_docx, size=settings.max_upload_bytes + 1)
        update, message = _upload(oversized)

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        assert "limit" in message.last_reply
        assert repository.get_all_resumes(42, settings=db) == []

    async def test_a_document_that_is_not_a_resume(
        self, db: Settings, settings: Settings, context: FakeContext, tmp_path: Path, jd_text: str
    ) -> None:
        posting = _write_docx(tmp_path / "posting.docx", jd_text * 3)
        update, message = _upload(FakeDocument(posting))

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        assert "does not look like a resume" in message.all_output
        assert repository.get_all_resumes(42, settings=db) == []

    async def test_a_too_short_document(
        self, db: Settings, context: FakeContext, tmp_path: Path
    ) -> None:
        stub = _write_docx(tmp_path / "stub.docx", "Jane Doe\nEXPERIENCE\nEngineer")
        update, message = _upload(FakeDocument(stub))

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        assert "too short to score" in message.all_output
        assert repository.get_all_resumes(42, settings=db) == []

    async def test_rejected_files_are_deleted_from_disk(
        self, db: Settings, settings: Settings, context: FakeContext, tmp_path: Path
    ) -> None:
        stub = _write_docx(tmp_path / "stub.docx", "Jane Doe\nEXPERIENCE\nEngineer")
        update, _ = _upload(FakeDocument(stub))

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        user_dir = settings.upload_dir / "42"
        leftovers = list(user_dir.glob("*")) if user_dir.is_dir() else []
        assert leftovers == []

    async def test_a_corrupt_document_is_reported_clearly(
        self, db: Settings, context: FakeContext, tmp_path: Path
    ) -> None:
        broken = tmp_path / "broken.docx"
        broken.write_bytes(b"not a real docx package")
        update, message = _upload(FakeDocument(broken))

        await handle_resume_upload(update, context)  # type: ignore[arg-type]

        assert "not a valid .docx" in message.all_output
        assert repository.get_all_resumes(42, settings=db) == []

    async def test_a_message_with_no_attachment(self, db: Settings, context: FakeContext) -> None:
        message = FakeMessage()
        await handle_resume_upload(FakeUpdate(message=message), context)  # type: ignore[arg-type]
        assert "could not find a file" in message.last_reply
