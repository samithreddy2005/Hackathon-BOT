"""DOCX text extraction via python-docx."""

from __future__ import annotations

import logging
from pathlib import Path

import docx
from docx.opc.exceptions import PackageNotFoundError

from ats_bot.errors import ParsingError
from ats_bot.parsing.cleaner import clean_text

__all__ = ["extract_text_from_docx"]

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str | Path) -> str:
    """Extract text from a DOCX file, including tables and headers/footers.

    Many resume templates put contact details in a header and skills in a table;
    reading paragraphs alone silently loses both, which then shows up as a bogus
    "email address not found" formatting penalty.

    Raises:
        ParsingError: If the file is missing or is not a valid DOCX package —
            most often a ``.doc`` file that was renamed rather than converted.
    """
    path = Path(file_path)
    if not path.is_file():
        raise ParsingError(f"DOCX file not found: {path.name}")

    parts: list[str] = []
    try:
        document = docx.Document(str(path))

        for section in document.sections:
            for container in (section.header, section.footer):
                for paragraph in container.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text.strip())

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                # A row's cells are deduplicated because merged cells repeat text.
                unique = list(dict.fromkeys(cells))
                if unique:
                    parts.append(" | ".join(unique))
    except PackageNotFoundError as exc:
        raise ParsingError(
            "This file is not a valid .docx document. If it is an older .doc file, "
            "open it in Word and save it as .docx or PDF first."
        ) from exc
    except Exception as exc:
        logger.exception("Failed to parse DOCX %s", path.name)
        raise ParsingError("This Word document could not be read. It may be corrupted.") from exc

    return clean_text("\n".join(parts))
